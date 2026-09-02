from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_txt_card(card_data):
    """Создаёт TXT файл с карточкой товара"""
    
    text = f"""{card_data.get('title', 'Без названия')}
{'=' * len(card_data.get('title', 'Без названия'))}

ОПИСАНИЕ:
{card_data.get('description', '')}

{'='*60}

КЛЮЧЕВЫЕ ПРЕИМУЩЕСТВА:
"""
    
    for i, adv in enumerate(card_data.get('advantages', [])[:7], 1):
        text += f"{i}. {adv}\n"
    
    text += f"\n{'='*60}\n\n"
    text += "ХАРАКТЕРИСТИКИ:\n"
    
    for char in card_data.get('characteristics', []):
        text += f"• {char}\n"
    
    text += f"\n{'='*60}\n\n"
    text += "ТЕКСТЫ ДЛЯ ИНФОГРАФИКИ:\n"
    
    for i, block in enumerate(card_data.get('infographic', [])[:5], 1):
        text += f"\nБлок {i}:\n{block}\n"
    
    keywords = card_data.get('keywords', [])
    if keywords:
        text += f"\n{'='*60}\n\n"
        text += "КЛЮЧЕВЫЕ СЛОВА:\n"
        text += ", ".join(keywords[:15])
    
    return text


def create_docx_card(card_data):
    """Создаёт DOCX файл (Word) с карточкой товара"""
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading(card_data.get('title', 'Карточка товара'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Описание
    doc.add_heading('Описание', level=1)
    doc.add_paragraph(card_data.get('description', ''))
    
    # Преимущества
    doc.add_heading('Ключевые преимущества', level=1)
    for adv in card_data.get('advantages', [])[:7]:
        doc.add_paragraph(f'• {adv}', style='List Bullet')
    
    # Характеристики
    doc.add_heading('Характеристики', level=1)
    for char in card_data.get('characteristics', []):
        doc.add_paragraph(f'• {char}', style='List Bullet')
    
    # Инфографика
    doc.add_heading('Тексты для инфографики', level=1)
    for i, block in enumerate(card_data.get('infographic', [])[:5], 1):
        doc.add_paragraph(f'Блок {i}: {block}')
    
    # Ключевые слова
    keywords = card_data.get('keywords', [])
    if keywords:
        doc.add_heading('Ключевые слова', level=1)
        doc.add_paragraph(', '.join(keywords[:15]))
    
    # Сохраняем в bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer