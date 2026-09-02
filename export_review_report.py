from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_review_report_txt(analysis, reviews):
    """Создаёт TXT отчёт об анализе отзывов"""
    
    summary = analysis.get('summary', {})
    
    text = f"""АНАЛИЗ ОТЗЫВОВ О ТОВАРЕ
{'='*50}

ОБЩАЯ СТАТИСТИКА:
Всего отзывов проанализировано: {summary.get('total_reviews', 0)}
Средний рейтинг: {summary.get('average_rating', 0)}/5
Положительных отзывов: {summary.get('positive_percentage', 0)}%
Отрицательных отзывов: {summary.get('negative_percentage', 0)}%

{'='*50}

ПРЕИМУЩЕСТВА ТОВАРА (по мнению покупателей):
"""
    
    for i, pro in enumerate(analysis.get('pros', [])[:5], 1):
        text += f"{i}. {pro}\n"
    
    text += f"\n{'='*50}\n\n"
    text += "НЕДОСТАТКИ ТОВАРА (по мнению покупателей):\n"
    
    for i, con in enumerate(analysis.get('cons', [])[:5], 1):
        text += f"{i}. {con}\n"
    
    text += f"\n{'='*50}\n\n"
    text += "БОЛЕВЫЕ ТОЧКИ КЛИЕНТОВ:\n"
    
    for i, pain in enumerate(analysis.get('customer_pain_points', [])[:5], 1):
        text += f"{i}. {pain}\n"
    
    text += f"\n{'='*50}\n\n"
    text += "РЕКОМЕНДАЦИИ ПО УЛУЧШЕНИЮ:\n"
    
    for i, rec in enumerate(analysis.get('recommendations', [])[:5], 1):
        text += f"{i}. {rec}\n"
    
    text += f"\n{'='*50}\n\n"
    text += f"ЦЕЛЕВАЯ АУДИТОРИЯ:\n{analysis.get('target_audience', 'Не определена')}\n"
    
    text += f"\n{'='*50}\n\n"
    text += f"КОНКУРЕНТНЫЕ ПРЕИМУЩЕСТВА:\n{analysis.get('competitor_advantages', 'Не указано')}\n"
    
    return text


def create_review_report_docx(analysis, reviews):
    """Создаёт DOCX отчёт об анализе отзывов"""
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Анализ отзывов о товаре', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    summary = analysis.get('summary', {})
    
    # Общая статистика
    doc.add_heading('Общая статистика', level=1)
    doc.add_paragraph(f"Всего отзывов: {summary.get('total_reviews', 0)}")
    doc.add_paragraph(f"Средний рейтинг: {summary.get('average_rating', 0)}/5")
    doc.add_paragraph(f"Положительных: {summary.get('positive_percentage', 0)}%")
    doc.add_paragraph(f"Отрицательных: {summary.get('negative_percentage', 0)}%")
    
    # Преимущества
    doc.add_heading('Преимущества товара', level=1)
    for pro in analysis.get('pros', [])[:5]:
        doc.add_paragraph(f"✅ {pro}", style='List Bullet')
    
    # Недостатки
    doc.add_heading('Недостатки товара', level=1)
    for con in analysis.get('cons', [])[:5]:
        doc.add_paragraph(f"❌ {con}", style='List Bullet')
    
    # Болевые точки
    doc.add_heading('Болевые точки клиентов', level=1)
    for pain in analysis.get('customer_pain_points', [])[:5]:
        doc.add_paragraph(f"⚠️ {pain}", style='List Bullet')
    
    # Рекомендации
    doc.add_heading('Рекомендации по улучшению', level=1)
    for rec in analysis.get('recommendations', [])[:5]:
        doc.add_paragraph(f"💡 {rec}", style='List Bullet')
    
    # Целевая аудитория
    doc.add_heading('Целевая аудитория', level=1)
    doc.add_paragraph(analysis.get('target_audience', 'Не определена'))
    
    # Конкурентные преимущества
    doc.add_heading('Конкурентные преимущества', level=1)
    doc.add_paragraph(analysis.get('competitor_advantages', 'Не указано'))
    
    # Сохраняем в bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer